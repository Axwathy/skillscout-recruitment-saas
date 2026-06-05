import io
import logging

import docx
import pdfplumber
from celery import shared_task
from django.utils import timezone

from apps.core.storage import download_file

from .models import ParsedResume, Resume
from .resume_parser import parse_resume_text

logger = logging.getLogger(__name__)


def extract_text_from_bytes(file_bytes: bytes, mime_type: str) -> str:
    file_obj = io.BytesIO(file_bytes)

    if mime_type == "application/pdf":
        with pdfplumber.open(file_obj) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n".join(pages)

    if mime_type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ]:
        doc = docx.Document(file_obj)
        return extract_docx_text(doc)

    logger.warning("Unsupported resume mime type: %s", mime_type)
    return ""


def extract_docx_text(doc: docx.Document) -> str:
    lines: list[str] = []

    def add_line(value: str) -> None:
        text = " ".join((value or "").split())
        if text:
            lines.append(text)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            add_line(paragraph.text)
        for paragraph in section.footer.paragraphs:
            add_line(paragraph.text)

    for paragraph in doc.paragraphs:
        add_line(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = " ".join(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )
                if cell_text and cell_text not in cells:
                    cells.append(cell_text)
            add_line(" | ".join(cells))

    return "\n".join(lines)


def extract_resume_text_from_bytes(resume: Resume, file_bytes: bytes) -> None:
    try:
        resume.status = Resume.Status.PROCESSING
        resume.save(update_fields=["status", "updated_at"])

        resume.raw_text = extract_text_from_bytes(file_bytes, resume.mime_type)
        resume.status = Resume.Status.PROCESSING
        resume.save(update_fields=["raw_text", "status", "updated_at"])

        parse_resume_with_llm.delay(str(resume.id))
    except Exception as exc:
        logger.error("Failed to extract text for resume %s: %s", resume.id, exc)
        resume.status = Resume.Status.ERROR
        resume.save(update_fields=["status", "updated_at"])


@shared_task
def extract_resume_text(resume_id: str):
    try:
        resume = Resume.objects.get(id=resume_id)
        resume.status = Resume.Status.PROCESSING
        resume.save(update_fields=["status", "updated_at"])

        # Fetch the file from Supabase to memory
        file_bytes = download_file("resumes", resume.file_url)
        extracted_text = extract_text_from_bytes(file_bytes, resume.mime_type)

        resume.raw_text = extracted_text
        resume.status = Resume.Status.PROCESSING
        resume.save(update_fields=["raw_text", "status", "updated_at"])

        try:
            parse_resume_with_llm.delay(str(resume.id))
        except Exception:
            # When Redis/Celery is unavailable in local development, keep the
            # pipeline usable by parsing inline after successful extraction.
            parse_resume_with_llm(str(resume.id))
        
    except Exception as e:
        logger.error(f"Failed to extract text for resume {resume_id}: {e}")
        try:
            resume = Resume.objects.get(id=resume_id)
            resume.status = Resume.Status.ERROR
            resume.save(update_fields=["status", "updated_at"])
        except Resume.DoesNotExist:
            pass


@shared_task(
    autoretry_for=(Exception,),
    retry_backoff=True,
    retry_jitter=True,
    retry_kwargs={"max_retries": 2},
)
def parse_resume_with_llm(resume_id: str):
    resume = Resume.objects.select_related("candidate", "application").get(id=resume_id)
    resume.status = Resume.Status.PROCESSING
    resume.save(update_fields=["status", "updated_at"])

    parsed_resume, _created = ParsedResume.objects.update_or_create(
        resume=resume,
        defaults={
            "candidate": resume.candidate,
            "application": resume.application,
            "status": ParsedResume.Status.PROCESSING,
        },
    )

    try:
        result = parse_resume_text(resume.raw_text)
    except Exception as exc:
        logger.error(f"Failed to parse resume {resume_id}: {exc}")
        parsed_resume.status = ParsedResume.Status.ERROR
        parsed_resume.validation_errors = [str(exc)]
        parsed_resume.parsed_at = timezone.now()
        parsed_resume.save(
            update_fields=["status", "validation_errors", "parsed_at", "updated_at"]
        )
        resume.status = Resume.Status.ERROR
        resume.save(update_fields=["status", "updated_at"])
        return

    parsed_resume.data = result.data
    parsed_resume.confidence = result.confidence
    parsed_resume.parser_model = result.model
    parsed_resume.validation_errors = result.validation_errors
    parsed_resume.token_usage = result.token_usage
    parsed_resume.estimated_cost = result.estimated_cost
    parsed_resume.status = ParsedResume.Status.COMPLETED
    parsed_resume.parsed_at = timezone.now()
    parsed_resume.save(
        update_fields=[
            "data",
            "confidence",
            "parser_model",
            "validation_errors",
            "token_usage",
            "estimated_cost",
            "status",
            "parsed_at",
            "updated_at",
        ]
    )

    resume.status = Resume.Status.COMPLETED
    resume.save(update_fields=["status", "updated_at"])

    try:
        from apps.ai_engine.tasks import generate_parsed_resume_embedding

        generate_parsed_resume_embedding.delay(str(parsed_resume.id))
    except Exception as exc:
        logger.warning(
            "Failed to enqueue embedding for parsed resume %s: %s",
            parsed_resume.id,
            exc,
        )
